from rest_framework.viewsets import GenericViewSet
from rest_framework import viewsets, mixins
from carsales.models import Dealer, DealerCenter, Car, Sale, Customer
from rest_framework.decorators import action
from rest_framework.response import Response
from django.db.models import Count, Avg, Max, Min
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from rest_framework.permissions import IsAuthenticated
from carsales.serializers import (
    DealerSerializer, DealerCenterSerializer,
    CarSerializer, SaleSerializer, CustomerSerializer
)
from rest_framework import serializers
from rest_framework.permissions import BasePermission

import openpyxl
from openpyxl.styles import Font
from django.http import HttpResponse
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework import status
from io import BytesIO
from openpyxl import Workbook
from docx import Document

# Базовый ViewSet для создания фильтрации по пользователю
class BaseUserViewSet(
        mixins.ListModelMixin,
        mixins.CreateModelMixin,
        mixins.UpdateModelMixin,
        mixins.RetrieveModelMixin,
        mixins.DestroyModelMixin,
        GenericViewSet
    ):

    def get_queryset(self):
        qs = super().get_queryset()
        user = self.request.user
        if user.is_superuser:
            # Если есть фильтрация по пользователю, то применяем её
            user_id = self.request.query_params.get("user_id")
            if user_id:
                qs = qs.filter(user_id=user_id)
        else:
            # Обычные пользователи видят только свои данные
            qs = qs.filter(user=user)
        return qs

# ViewSet для модели Dealer
class DealerViewSet(BaseUserViewSet):
    queryset = Dealer.objects.all()
    serializer_class = DealerSerializer
    permission_classes = [IsAuthenticated]

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, *args, **kwargs):
        stats = Dealer.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )

        serializer = self.StatsSerializer(instance=stats)
        return Response(serializer.data)

# ViewSet для модели DealerCenter
class DealerCenterViewSet(BaseUserViewSet):
    queryset = DealerCenter.objects.all()
    serializer_class = DealerCenterSerializer

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, *args, **kwargs):
        stats = DealerCenter.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )

        serializer = self.StatsSerializer(instance=stats)
        return Response(serializer.data)

# ViewSet для модели Car
class CarViewSet(BaseUserViewSet):
    queryset = Car.objects.all()
    serializer_class = CarSerializer

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, *args, **kwargs):
        stats = Car.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )

        serializer = self.StatsSerializer(instance=stats)
        return Response(serializer.data)


    @action(detail=False, methods=["GET"], url_path="export-excel")
    def export_to_excel(self, request, *args, **kwargs):
        # Создаем Excel-файл
        workbook = Workbook()  # Теперь импорт этого класса есть
        sheet = workbook.active
        sheet.title = "Cars"

        # Добавляем заголовки
        headers = ["ID", "Dealer", "Model", "Year", "Price", "Dealer Center"]
        for col_num, header in enumerate(headers, start=1):
            sheet.cell(row=1, column=col_num, value=header)

        # Добавляем данные
        cars = self.queryset.select_related("dealer_FK", "dealer_center_FK")
        for row_num, car in enumerate(cars, start=2):
            sheet.cell(row=row_num, column=1, value=car.id)
            sheet.cell(row=row_num, column=2, value=car.dealer_FK.name if car.dealer_FK else "")
            sheet.cell(row=row_num, column=3, value=car.car_model)
            sheet.cell(row=row_num, column=4, value=car.year)
            sheet.cell(row=row_num, column=5, value=car.price)
            sheet.cell(row=row_num, column=6, value=car.dealer_center_FK.headquarters_location if car.dealer_center_FK else "")

        # Сохранение и отправка файла
        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = 'attachment; filename="cars.xlsx"'
        return response

    @action(detail=False, methods=["GET"], url_path="export-word")
    def export_to_word(self, request, *args, **kwargs):
        # Создаем Word-документ
        document = Document()
        document.add_heading("Car List", level=1)

        # Добавляем данные в таблицу
        cars = self.queryset.select_related("dealer_FK", "dealer_center_FK")
        table = document.add_table(rows=1, cols=6)
        headers = ["ID", "Dealer", "Model", "Year", "Price", "Dealer Center"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header

        for car in cars:
            row = table.add_row().cells
            row[0].text = str(car.id)
            row[1].text = car.dealer_FK.name if car.dealer_FK else ""
            row[2].text = car.car_model
            row[3].text = car.year
            row[4].text = car.price
            row[5].text = car.dealer_center_FK.headquarters_location if car.dealer_center_FK else ""

        # Сохранение и отправка файла
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = HttpResponse(
            buffer,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = 'attachment; filename="cars.docx"'
        return response
        
# ViewSet для модели Customer
class CustomerViewSet(BaseUserViewSet):
    queryset = Customer.objects.all()
    serializer_class = CustomerSerializer

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, *args, **kwargs):
        stats = Customer.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )

        serializer = self.StatsSerializer(instance=stats)
        return Response(serializer.data)

# ViewSet для модели Sale
class SaleViewSet(BaseUserViewSet):
    queryset = Sale.objects.all()
    serializer_class = SaleSerializer

    class StatsSerializer(serializers.Serializer):
        count = serializers.IntegerField()
        avg = serializers.FloatField()
        max = serializers.IntegerField()
        min = serializers.IntegerField()

    @action(detail=False, methods=["GET"], url_path="stats")
    def get_stats(self, *args, **kwargs):
        stats = Sale.objects.aggregate(
            count=Count("*"),
            avg=Avg("id"),
            max=Max("id"),
            min=Min("id"),
        )

        serializer = self.StatsSerializer(instance=stats)
        return Response(serializer.data)


# ViewSet для User
class UserViewSet(GenericViewSet):
    @action(url_path="info", methods=["GET"], detail=False)
    def get_info(self, request, *args, **kwargs):
        data = {
            "is_authenticated": request.user.is_authenticated
        }
        if request.user.is_authenticated:
            data.update({
                "username": request.user.username,
                "user_id": request.user.id,
                "is_superuser": request.user.is_superuser
            })
        return Response(data)

    @action(detail=False, methods=["GET"], url_path="list-users")
    def list_users(self, request, *args, **kwargs):
        if not request.user.is_superuser:
            return Response({"detail": "Permission denied."}, status=403)
        
        users = User.objects.values("id", "username")
        return Response(users)

    @action(url_path="login", methods=["POST"], detail=False)
    def login(self, request, *args, **kwargs):
        username = request.data.get("user")
        password = request.data.get("password")

        user = authenticate(request, username=username, password=password)
        if user:
            login(request, user)
        return Response({})

    @action(url_path="logout", methods=["POST"], detail=False)
    def logout(self, request, *args, **kwargs):
        logout(request)
        return Response({})



class UserProfileViewSet(viewsets.GenericViewSet):
    permission_classes = [IsAuthenticated]

    class OTPSerializer(serializers.Serializer):
        key = serializers.CharField()

    class OTPRequired(BasePermission):
        def has_permission(self, request, view):
            return bool(request.user and cache.get('otp_good', False))

    @action(detail=False, url_path="check-login", methods=['GET'], permission_classes=[])
    def get_check_login(self, request, *args, **kwargs):
        return Response({
            'is_authenticated': self.request.user.is_authenticated
        })
    
    @action(detail=False, url_path="login", methods=['GET'], permission_classes=[])
    def use_login(self, request, *args, **kwargs):
        user = authenticate(username='username', password='pass')
        if user:
            login(request, user)
        return Response({
            'is_authenticated': bool(user)
        })

    @action(detail=False, url_path='otp-login', methods=['POST'], serializer_class=OTPSerializer)
    def otp_login(self, *args, **kwargs):
        totp = pyotp.TOTP(self.request.user.userprofile.opt_key)
        
        serializer = self.get_serializer(data=self.request.data)
        serializer.is_valid(raise_exception=True)

        success = False
        if totp.now() == serializer.validated_data['key']:
            cache.set('otp_good', True, 600)  # Устанавливаем время жизни флага на 10 минут
            success = True

        return Response({
            'success': success
        })
    
    @action(detail=False, url_path='otp-status')
    def get_otp_status(self, *args, **kwargs):
        otp_good = cache.get('otp_good', False)
        return Response({
            'otp_good': otp_good
        })
    
    @action(detail=False, url_path='otp-required', permission_classes=[OTPRequired])
    def page_with_otp_required(self, *args, **kwargs):
        return Response({
            'success': True
        })

    @action(detail=False, url_path="edit-object", methods=['POST'], permission_classes=[OTPRequired])
    def edit_object(self, request, *args, **kwargs):
        # Логика редактирования объекта
        return Response({
            'success': True
        })
